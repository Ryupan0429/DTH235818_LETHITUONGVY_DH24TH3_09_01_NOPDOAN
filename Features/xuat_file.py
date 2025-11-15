import tkinter as tk
from tkinter import messagebox, filedialog
from db import get_connection
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    messagebox.showerror("Thiếu thư viện", 
        "Không tìm thấy thư viện 'python-docx'.\n"
        "Vui lòng chạy: pip install python-docx")
    raise

# --- THÔNG TIN CỬA HÀNG (Bạn có thể sửa tại đây) ---
TEN_CUA_HANG = "CỬA HÀNG NÔNG DƯỢC TƯỜNG VY"
DIA_CHI_CUA_HANG = "123, Đường ABC, Phường XYZ, TP. Long Xuyên, An Giang"
SDT_CUA_HANG = "0901 029 661"
# -----------------------------------------------

def _set_table_font(table):
    """Helper: Thiết lập font chữ chung cho bảng."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

def _add_paragraph(doc, text, style=None, size=None, bold=False, align=None):
    """Helper: Thêm một đoạn văn bản."""
    p = doc.add_paragraph(text, style=style)
    if size:
        p.runs[0].font.size = Pt(size)
    if bold:
        p.runs[0].font.bold = True
    if align:
        p.alignment = align
    p.runs[0].font.name = 'Times New Roman'
    return p

def _format_currency(value):
    """Helper: Định dạng số sang tiền tệ."""
    try:
        return f"{int(value):,.0f} VNĐ"
    except (ValueError, TypeError):
        return "0 VNĐ"

def export_invoice_to_word(parent_window, mahd):
    """Truy vấn CSDL và xuất Hóa Đơn (HD) ra file .docx."""
    conn = None
    try:
        # 1. Lấy đường dẫn lưu file
        save_path = filedialog.asksaveasfilename(
            parent=parent_window,
            title=f"Lưu Hóa đơn {mahd}",
            initialfile=f"HoaDon_{mahd}.docx",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if not save_path:
            return # Người dùng nhấn Hủy

        # 2. Kết nối và truy vấn dữ liệu
        conn = get_connection()
        cur = conn.cursor()
        
        # Lấy thông tin Hóa đơn và Khách hàng
        sql_header = """
        SELECT H.NgayGD, K.TenKH, K.SDT, K.QueQuan, H.TongGT
        FROM dbo.HoaDon H
        JOIN dbo.KhachHang K ON H.MaKH = K.MaKH
        WHERE H.MaHD = ?
        """
        cur.execute(sql_header, (mahd,))
        header = cur.fetchone()
        
        if not header:
            messagebox.showerror("Lỗi", "Không tìm thấy thông tin hóa đơn.", parent=parent_window)
            return

        # Lấy chi tiết Hóa đơn
        sql_details = "SELECT MaSP, TenSP, SoLuong, DVTinh, DonGia, ThanhTien FROM dbo.ChiTietHoaDon WHERE MaHD = ?"
        cur.execute(sql_details, (mahd,))
        details = cur.fetchall()
        conn.close()

        # 3. Tạo tài liệu Word
        doc = Document()
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(11)

        # Thông tin cửa hàng
        _add_paragraph(doc, TEN_CUA_HANG, size=14, bold=True)
        _add_paragraph(doc, f"Địa chỉ: {DIA_CHI_CUA_HANG}")
        _add_paragraph(doc, f"SĐT: {SDT_CUA_HANG}")
        
        # Tiêu đề
        _add_paragraph(doc, "\nHÓA ĐƠN BÁN HÀNG", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_paragraph(doc, f"Mã HĐ: {mahd}", align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_paragraph(doc, f"Ngày: {header.NgayGD.strftime('%d/%m/%Y')}\n", align=WD_ALIGN_PARAGRAPH.CENTER)

        # Thông tin khách hàng
        _add_paragraph(doc, "Thông tin Khách hàng:", bold=True)
        _add_paragraph(doc, f"Họ tên: {header.TenKH}")
        _add_paragraph(doc, f"SĐT: {header.SDT}")
        _add_paragraph(doc, f"Địa chỉ: {header.QueQuan}\n")

        # Bảng chi tiết
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Mã SP'
        hdr_cells[1].text = 'Tên Sản Phẩm'
        hdr_cells[2].text = 'SL'
        hdr_cells[3].text = 'ĐVT'
        hdr_cells[4].text = 'Đơn Giá'
        hdr_cells[5].text = 'Thành Tiền'

        for masp, tensp, sl, dvt, dongia, thanhtien in details:
            row_cells = table.add_row().cells
            row_cells[0].text = masp
            row_cells[1].text = tensp
            row_cells[2].text = str(sl)
            row_cells[3].text = dvt
            row_cells[4].text = _format_currency(dongia)
            row_cells[5].text = _format_currency(thanhtien)
        
        _set_table_font(table)
        
        # Tổng cộng
        _add_paragraph(doc, f"\nTổng cộng: {_format_currency(header.TongGT)}", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _add_paragraph(doc, "\n(Đã bao gồm thuế GTGT nếu có)", style='Caption', align=WD_ALIGN_PARAGRAPH.RIGHT)

        # 4. Lưu file
        doc.save(save_path)
        messagebox.showinfo("Thành công", f"Đã xuất hóa đơn thành công!\n{save_path}", parent=parent_window)

    except Exception as e:
        messagebox.showerror("Lỗi Xuất file", f"Không thể xuất file Hóa đơn:\n{e}", parent=parent_window)
    finally:
        if conn: conn.close()

def export_import_bill_to_word(parent_window, sopn):
    """Truy vấn CSDL và xuất Phiếu Nhập (PN) ra file .docx."""
    conn = None
    try:
        # 1. Lấy đường dẫn lưu file
        save_path = filedialog.asksaveasfilename(
            parent=parent_window,
            title=f"Lưu Phiếu Nhập {sopn}",
            initialfile=f"PhieuNhap_{sopn}.docx",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if not save_path:
            return

        # 2. Kết nối và truy vấn dữ liệu
        conn = get_connection()
        cur = conn.cursor()
        
        # Lấy thông tin Phiếu Nhập
        sql_header = "SELECT NgayNhap, NguoiNhap, NguonNhap, TongGGT FROM dbo.PhieuNhap WHERE SoPN = ?"
        cur.execute(sql_header, (sopn,))
        header = cur.fetchone()
        
        if not header:
            messagebox.showerror("Lỗi", "Không tìm thấy thông tin phiếu nhập.", parent=parent_window)
            return

        # Lấy chi tiết Phiếu Nhập
        sql_details = "SELECT MaSP, TenSP, SoLuong, DVTinh, DonGia, ThanhTien FROM dbo.ChiTietPhieuNhap WHERE SoPN = ?"
        cur.execute(sql_details, (sopn,))
        details = cur.fetchall()
        conn.close()

        # 3. Tạo tài liệu Word
        doc = Document()
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(11)

        # Thông tin cửa hàng
        _add_paragraph(doc, TEN_CUA_HANG, size=14, bold=True)
        _add_paragraph(doc, f"Địa chỉ: {DIA_CHI_CUA_HANG}")
        
        # Tiêu đề
        _add_paragraph(doc, "\nPHIẾU NHẬP KHO", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_paragraph(doc, f"Số PN: {sopn}", align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_paragraph(doc, f"Ngày nhập: {header.NgayNhap.strftime('%d/%m/%Y')}\n", align=WD_ALIGN_PARAGRAPH.CENTER)

        # Thông tin phiếu
        _add_paragraph(doc, f"Người nhập: {header.NguoiNhap}")
        _add_paragraph(doc, f"Nguồn nhập (NCC): {header.NguonNhap}\n")

        # Bảng chi tiết
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Mã SP'
        hdr_cells[1].text = 'Tên Sản Phẩm'
        hdr_cells[2].text = 'SL'
        hdr_cells[3].text = 'ĐVT'
        hdr_cells[4].text = 'Đơn Giá (Nhập)'
        hdr_cells[5].text = 'Thành Tiền'

        for masp, tensp, sl, dvt, dongia, thanhtien in details:
            row_cells = table.add_row().cells
            row_cells[0].text = masp
            row_cells[1].text = tensp
            row_cells[2].text = str(sl)
            row_cells[3].text = dvt
            row_cells[4].text = _format_currency(dongia)
            row_cells[5].text = _format_currency(thanhtien)
        
        _set_table_font(table)
        
        # Tổng cộng
        _add_paragraph(doc, f"\nTổng giá trị nhập: {_format_currency(header.TongGGT)}", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        
        # 4. Lưu file
        doc.save(save_path)
        messagebox.showinfo("Thành công", f"Đã xuất phiếu nhập thành công!\n{save_path}", parent=parent_window)

    except Exception as e:
        messagebox.showerror("Lỗi Xuất file", f"Không thể xuất file Phiếu nhập:\n{e}", parent=parent_window)
    finally:
        if conn: conn.close()